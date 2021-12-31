import docx
import datetime
import os

root = '/Users/echooo/Documents/中科院/周报'
module = '/Users/echooo/Documents/中科院/周报/周报模版.docx'
today = datetime.date.today()
monday = today + datetime.timedelta(days=-4)
next_monday = monday + datetime.timedelta(days=7)
next_friday = today + datetime.timedelta(days=7)
last_monday = monday + datetime.timedelta(days=-7)
last_friday = today + datetime.timedelta(days=-7)

str_mon = f'{monday.month}月{monday.day}日'
str_today = f'{today.month}月{today.day}日'
str_next_mon = f'{next_monday.month}月{next_monday.day}日'
str_next_fri = f'{next_friday.month}月{next_friday.day}日'

last_report = f'/Users/echooo/Documents/中科院/周报/尹鹏宇{last_monday.month}.{last_monday.day}-{last_friday.day}.docx'
document = docx.Document(module)
paragraphs = document.paragraphs
for par in paragraphs:
    print(par.text)

tables = document.tables
date = tables[0].cell(0,1).paragraphs[0]
for run in date.runs:
    if 'monday' in run.text:
        run.text = run.text.replace('monday', str_mon)
    if 'friday' in run.text:
        run.text = run.text.replace('friday', str_today)

next_date = tables[1].cell(0,1).paragraphs[0].add_run(f'{str_next_mon}至{str_next_fri}工作计划')


###########本周工作计划##############
last_report_document = docx.Document(last_report)
last_tables = last_report_document.tables
for i, par in enumerate(last_tables[1].cell(1,1).paragraphs):
    par = par.text
    print(par)
    this_week_work_plan = tables[0].cell(1,1).paragraphs[0].add_run(f'{i+1}.{par}\n')
# completion_summary = tables[0].cell(1,2).paragraphs[0].add_run('')

###########下周工作计划##############
# next_week_work_plan = tables[1].cell(1,1).paragraphs[0].add_run('1. 优化船只检测算法代码\n2. 实现styleGAN算法\n3. 筹备论文工作')
print(today)
print(monday)

file_name = f'尹鹏宇{monday.month}.{monday.day}-{today.day}.docx'
document.save(os.path.join(root, file_name))